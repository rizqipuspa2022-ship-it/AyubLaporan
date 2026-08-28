import streamlit as st
from pathlib import Path
from io import BytesIO
import os, re, json, statistics
from docx import Document
from pypdf import PdfReader
from openpyxl import load_workbook
import xlsxwriter

APP_DIR = Path(__file__).parent
SCHEMA = json.loads((APP_DIR / "schema.json").read_text(encoding="utf-8"))
ALL_FIELDS = [x for x in SCHEMA["legacy_mailmerge_fields"] if x]
CORE_FIELDS = SCHEMA["core_fields"]

st.set_page_config(page_title="Ayub Laporan", page_icon="📑", layout="wide")
st.title("📑 Ayub Laporan")
st.caption("Upload dokumennya, bukan datanya. Sistem membaca PDF/Word/Excel, menghitung statistik, lalu membentuk source mailmerge + evidence log + Word laporan.")

# ---------- helpers ----------
def blank_fields():
    return {k: "" for k in ALL_FIELDS}

def put(fields, evidence, key, value, source, confidence=0.90, note=""):
    if key not in fields or value in [None, ""]:
        return
    value = str(value).strip()
    if not value:
        return
    if not fields.get(key):
        fields[key] = value
        evidence[key] = {"source": source, "confidence": confidence, "note": note}

def fmt_num(v, digits=2):
    if v is None:
        return ""
    try:
        f = float(v)
        if abs(f - round(f)) < 1e-9:
            return str(int(round(f)))
        return f"{f:.{digits}f}".rstrip("0").rstrip(".")
    except Exception:
        return str(v)

def extract_pdf(data):
    return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)

def extract_docx(data):
    d = Document(BytesIO(data)); out=[]
    out += [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            vals=[c.text.strip() for c in r.cells]
            if any(vals): out.append(" | ".join(vals))
    return "\n".join(out)

def extract_xlsx_preview(data, max_rows=120):
    wb=load_workbook(BytesIO(data), data_only=True, read_only=True)
    out=[]
    for ws in wb.worksheets:
        out.append(f"### SHEET {ws.title}")
        for ridx,row in enumerate(ws.iter_rows(values_only=True),1):
            if ridx>max_rows: break
            vals=[str(v).strip() for v in row if v not in (None,"")]
            if vals: out.append(" | ".join(vals))
    return "\n".join(out)

def extract_uploaded(u):
    ext=Path(u.name).suffix.lower(); b=u.getvalue()
    if ext==".pdf": return extract_pdf(b)
    if ext==".docx": return extract_docx(b)
    if ext==".xlsx": return extract_xlsx_preview(b)
    if ext in [".csv",".txt"]: return b.decode("utf-8",errors="ignore")
    return ""

def normalize_number(s):
    return str(s).replace(",", ".").strip()

def norm_header(v):
    return re.sub(r"[^A-Z0-9]+", " ", str(v or "").upper()).strip()

def find_header_row(ws, required_groups, scan_rows=20):
    """required_groups = list of sets; each set is a group of acceptable header aliases."""
    for r in range(1, min(ws.max_row, scan_rows)+1):
        vals=[norm_header(ws.cell(r,c).value) for c in range(1, ws.max_column+1)]
        ok=True
        for aliases in required_groups:
            if not any(any(alias in cell for alias in aliases) for cell in vals):
                ok=False; break
        if ok:
            return r
    return None

def header_map(ws, row):
    result={}
    for c in range(1, ws.max_column+1):
        h=norm_header(ws.cell(row,c).value)
        if h and h not in result: result[h]=c
    return result

def find_col(hmap, aliases):
    for h,c in hmap.items():
        if any(a in h for a in aliases): return c
    return None

def parse_participant_workbook(data, name, fields, ev):
    wb=load_workbook(BytesIO(data), data_only=True, read_only=True)
    for ws in wb.worksheets:
        hr=find_header_row(ws,[{"NO"},{"NAMA"},{"GENDER"},{"PENDIDIKAN"}])
        if not hr: continue
        hm=header_map(ws,hr)
        c_no=find_col(hm,["NO"]); c_name=find_col(hm,["NAMA"])
        c_gender=find_col(hm,["GENDER","JENIS KELAMIN"]); c_edu=find_col(hm,["PENDIDIKAN"])
        rows=[]
        for r in range(hr+1, ws.max_row+1):
            no=ws.cell(r,c_no).value if c_no else None
            nm=ws.cell(r,c_name).value if c_name else None
            if isinstance(no,(int,float)) and nm not in (None,""):
                rows.append(r)
        if not rows: continue
        put(fields,ev,"realisasi_peserta",len(rows),name,0.99,"Dihitung dari baris peserta pada workbook")
        if c_gender:
            male=female=0
            for r in rows:
                g=norm_header(ws.cell(r,c_gender).value)
                if g in {"L","LAKI LAKI","MALE"}: male+=1
                elif g in {"P","PEREMPUAN","FEMALE"}: female+=1
            put(fields,ev,"realisasi_laki",male,name,0.99,"Dihitung dari kolom GENDER")
            put(fields,ev,"realisasi_perempuan",female,name,0.99,"Dihitung dari kolom GENDER")
        if c_edu:
            counts={"S3":0,"S2":0,"D4_S1":0,"D3":0,"D1":0,"SMA_sederajat":0,"Unknown_School":0}
            for r in rows:
                raw=norm_header(ws.cell(r,c_edu).value)
                compact=raw.replace(" ","")
                if compact in {"S3","DOKTOR","DOCTORAL"}: counts["S3"]+=1
                elif compact in {"S2","MAGISTER","MASTER"}: counts["S2"]+=1
                elif compact in {"S1","D4","DIV","SARJANA"}: counts["D4_S1"]+=1
                elif compact in {"D3","DIII"}: counts["D3"]+=1
                elif compact in {"D1","DI"}: counts["D1"]+=1
                elif any(x in compact for x in ["SMA","SMK","SLTA"]): counts["SMA_sederajat"]+=1
                else: counts["Unknown_School"]+=1
            for k,v in counts.items(): put(fields,ev,k,v,name,0.99,"Dihitung dari kolom PENDIDIKAN")
        return True
    return False

def parse_prepost_workbook(data, name, fields, ev):
    wb=load_workbook(BytesIO(data), data_only=True, read_only=True)
    for ws in wb.worksheets:
        hr=find_header_row(ws,[{"NO"},{"NAMA"},{"PRETEST","PRE TEST"},{"POSTTEST","POST TEST"}])
        if not hr: continue
        hm=header_map(ws,hr)
        c_no=find_col(hm,["NO"]); c_name=find_col(hm,["NAMA"])
        c_pre=find_col(hm,["PRETEST","PRE TEST"]); c_post=find_col(hm,["POSTTEST","POST TEST"])
        if not c_pre or not c_post: continue
        pres=[]; posts=[]
        for r in range(hr+1, ws.max_row+1):
            no=ws.cell(r,c_no).value if c_no else None
            nm=ws.cell(r,c_name).value if c_name else None
            pre=ws.cell(r,c_pre).value; post=ws.cell(r,c_post).value
            if isinstance(no,(int,float)) and nm not in (None,"") and isinstance(pre,(int,float)) and isinstance(post,(int,float)):
                pres.append(float(pre)); posts.append(float(post))
        if not pres: continue
        put(fields,ev,"Jenis_Ujian","Pretest dan Posttest",name,0.99)
        put(fields,ev,"peserta_ujian",len(pres),name,0.99,"Dihitung dari peserta dengan nilai pretest dan posttest")
        put(fields,ev,"Min_Kompre_Pretest",fmt_num(min(pres)),name,0.99)
        put(fields,ev,"Max_Kompre_Pretest",fmt_num(max(pres)),name,0.99)
        put(fields,ev,"Rata2_Kompre_Pretest",fmt_num(statistics.mean(pres)),name,0.99)
        put(fields,ev,"Min_NA_Posttest",fmt_num(min(posts)),name,0.99)
        put(fields,ev,"Max_NA_Posttest",fmt_num(max(posts)),name,0.99)
        put(fields,ev,"Rata2_NA_Posttest",fmt_num(statistics.mean(posts)),name,0.99)
        put(fields,ev,"Selisih_Pretest_Posttest",fmt_num(statistics.mean(posts)-statistics.mean(pres)),name,0.99,"Rata-rata posttest dikurangi rata-rata pretest")
        return True
    return False

def parse_structured_xlsx(docs, fields, ev):
    for d in docs:
        if d.get("ext") != ".xlsx" or not d.get("bytes"): continue
        data=d["bytes"]; name=d["name"]
        try:
            parse_participant_workbook(data,name,fields,ev)
        except Exception:
            pass
        try:
            parse_prepost_workbook(data,name,fields,ev)
        except Exception:
            pass

def deterministic_extract(docs):
    fields=blank_fields(); ev={}
    full="\n".join(x["text"] for x in docs)
    upper=full.upper()

    # Nama program dan angkatan
    m=re.search(r"PELATIHAN JARAK JAUH\s+([^\n]+(?:\n[^\n]+){0,2}?)\s+ANGKATAN\s+([IVXLCDM]+)", full, re.I)
    if m:
        title="Pelatihan Jarak Jauh "+re.sub(r"\s+"," ",m.group(1)).strip()
        title=re.sub(r"TAHUN ANGGARAN.*$","",title,flags=re.I).strip()
        put(fields,ev,"nama_pelatihan",title,"dokumen sumber",0.93)
        put(fields,ev,"akt",m.group(2),"dokumen sumber",0.93)

    # Pengumuman
    m=re.search(r"NOMOR\s+(PENG[-–][A-Z0-9./-]+)", full, re.I)
    if m: put(fields,ev,"no_pengumuman",m.group(1),"pengumuman",0.99)
    m=re.search(r"pada tanggal\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})", full, re.I)
    if m: put(fields,ev,"tgl_pengumuman",m.group(1),"pengumuman",0.95)

    # Periode: mendukung '17 s.d. 21 Oktober 2022' dan '17 Oktober 2022 sampai dengan 21 Oktober 2022'
    m=re.search(r"(\d{1,2})\s+(?:s\.d\.|s/d|sampai dengan)\s+(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})",full,re.I)
    if m:
        start=f"{m.group(1)} {m.group(3)} {m.group(4)}"; end=f"{m.group(2)} {m.group(3)} {m.group(4)}"
        put(fields,ev,"tgl_mulai",start,"jadwal/pengumuman",0.95); put(fields,ev,"tgl_selesai",end,"jadwal/pengumuman",0.95)
    else:
        m=re.search(r"(\d{1,2}\s+[A-Za-z]+\s+\d{4})\s+sampai dengan\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})",full,re.I)
        if m:
            put(fields,ev,"tgl_mulai",m.group(1),"rekap kehadiran",0.95); put(fields,ev,"tgl_selesai",m.group(2),"rekap kehadiran",0.95)
    if fields.get("tgl_mulai"):
        put(fields,ev,"non_tatap_muka_(mulai)",fields["tgl_mulai"],"periode PJJ",0.95)
    if fields.get("tgl_selesai"):
        put(fields,ev,"non_tatap_muka_(selesai)",fields["tgl_selesai"],"periode PJJ",0.95)

    if "PELATIHAN JARAK JAUH" in upper:
        put(fields,ev,"metode","Pelatihan Jarak Jauh","KAP/pengumuman",0.99)
    if "PUSDIKLAT PAJAK" in upper:
        put(fields,ev,"unit_asal","Direktorat Jenderal Pajak","dokumen peserta/KAP",0.85)
        put(fields,ev,"tempat_pelatihan","Pusdiklat Pajak","dokumen evaluasi",0.90)

    # PIC / Manajer kelas
    m=re.search(r"Manajer\s+Kelas\s*:?\s*([^\n|]+)",full,re.I)
    if m:
        pic=re.sub(r"\s+"," ",m.group(1)).strip(" :-")
        put(fields,ev,"PIC",pic,"rekap kehadiran",0.99,"Diambil dari label Manajer Kelas")

    # Jumlah peserta dari rekomendasi
    m=re.search(r"Jumlah Peserta Pembelajaran\s*:?\s*(\d+)\s*Orang",full,re.I)
    if m: put(fields,ev,"realisasi_peserta",m.group(1),"rekomendasi perbaikan",0.99)

    # Evaluasi penyelenggaraan dan pengajar
    m=re.search(r"Secara umum pembelajaran ini mendapatkan nilai rata-rata kinerja sebesar\s*([0-9]+[,.][0-9]+)",full,re.I)
    if not m: m=re.search(r"nilai rata-rata kinerja sebesar\s*([0-9]+[,.][0-9]+)",full,re.I)
    if m:
        val=normalize_number(m.group(1)); put(fields,ev,"indeks_rata2_evagara_PJJ",val,"rekomendasi perbaikan",0.99)
        put(fields,ev,"kriteria_idx_evagara_PJJ","sangat baik","rekomendasi perbaikan",0.99)
    m=re.search(r"evaluasi pengajar[\s\S]{0,500}?nilai rata-rata kinerja[\s\S]{0,250}?sebesar\s*([0-9]+[,.][0-9]+)",full,re.I)
    if m:
        val=normalize_number(m.group(1)); put(fields,ev,"indeks_kemampuan_mengajar",val,"rekomendasi perbaikan",0.99)
        put(fields,ev,"indeks_rata2_evajar",val,"rekomendasi perbaikan",0.99)
        put(fields,ev,"kriteria_idx_evajar","sangat baik","rekomendasi perbaikan",0.99)

    rec=[]
    if re.search(r"Ketercukupan waktu penyelenggaraan",full,re.I): rec.append("Ketercukupan waktu penyelenggaraan PJJ dengan jumlah materi yang diberikan")
    if re.search(r"Ketercukupan waktu dalam mengerjakan penugasan",full,re.I): rec.append("Ketercukupan waktu dalam mengerjakan penugasan, kuis, atau ujian")
    if rec: put(fields,ev,"rec_penyelenggaraan_pelatihan","; ".join(rec),"rekomendasi perbaikan",0.99)

    # KAP
    m=re.search(r"TUJUAN PROGRAM\s+([\s\S]{20,1200}?)(?:KEBUTUHAN STRATEGIS|SASARAN)",full,re.I)
    if m: put(fields,ev,"maksud_program",re.sub(r"\s+"," ",m.group(1)).strip(),"KAP",0.95)
    m=re.search(r"SASARAN \(TARGET LEARNERS\)\s+([\s\S]{10,700}?)(?:MODEL PEMBELAJARAN|STANDAR KOMPETENSI)",full,re.I)
    if m: put(fields,ev,"diharapkan_mampu",re.sub(r"\s+"," ",m.group(1)).strip(),"KAP",0.85,"Ditarik dari sasaran program")

    # Excel dihitung secara terstruktur, bukan dicari sebagai teks
    parse_structured_xlsx(docs,fields,ev)

    if fields.get("nama_pelatihan"):
        fields["nama_pelatihan_upper"]=fields["nama_pelatihan"].upper()
        fields["nama_pelatihan_lengkap"]=fields["nama_pelatihan"]+(f" Angkatan {fields.get('akt')}" if fields.get('akt') else "")
    return fields, ev

def get_api_key():
    try: return st.secrets.get("OPENAI_API_KEY","") or os.getenv("OPENAI_API_KEY","") or st.session_state.get("api_key","")
    except: return os.getenv("OPENAI_API_KEY","") or st.session_state.get("api_key","")

def ai_enrich(fields, evidence, docs, model):
    from openai import OpenAI
    client=OpenAI(api_key=get_api_key())
    missing=[k for k in ALL_FIELDS if not fields.get(k)]
    payload="\n\n".join(f"===== {d['name']} =====\n{d['text']}" for d in docs)[:180000]
    prompt=f'''Ekstrak data laporan penyelenggaraan dari dokumen berikut. Jangan mengarang. Isi hanya field yang benar-benar didukung sumber. Untuk PJJ gunakan field berakhiran _PJJ, bukan _elearning. Jangan menimpa hasil hitung Excel. Return JSON saja: {{"fields":{{}},"evidence":{{"field":{{"source":"nama file","confidence":0.0,"note":"bukti singkat"}}}}}}. Field yang perlu dicari: {json.dumps(missing,ensure_ascii=False)}\n\n{payload}'''
    r=client.responses.create(model=model,input=prompt)
    raw=re.sub(r"^```(?:json)?\s*|\s*```$","",r.output_text.strip())
    obj=json.loads(raw)
    for k,v in obj.get("fields",{}).items():
        if k in fields and not fields.get(k) and v not in [None,""]:
            fields[k]=str(v); evidence[k]=obj.get("evidence",{}).get(k,{"source":"AI","confidence":0.75,"note":"AI extraction"})
    return fields,evidence

def replace_p(p,mapping):
    text="".join(r.text for r in p.runs)
    if not text:return
    new=text
    for k,v in mapping.items(): new=new.replace(f"<<{k}>>",str(v or "")).replace(f"[@{k}]",str(v or ""))
    if new!=text:
        if p.runs:
            p.runs[0].text=new
            for r in p.runs[1:]:r.text=""
        else:p.text=new

def build_docx(mapping, template_bytes=None):
    d=Document(BytesIO(template_bytes)) if template_bytes else Document()
    if template_bytes:
        for p in d.paragraphs: replace_p(p,mapping)
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs: replace_p(p,mapping)
    else:
        d.add_heading("Laporan Penyelenggaraan",0)
        d.add_paragraph(mapping.get("nama_pelatihan_lengkap") or mapping.get("nama_pelatihan") or "Kegiatan Pembelajaran")
        for k in CORE_FIELDS:
            if mapping.get(k): d.add_paragraph(f"{k}: {mapping[k]}")
    out=BytesIO(); d.save(out); out.seek(0); return out

def build_xlsx(mapping,evidence):
    out=BytesIO(); wb=xlsxwriter.Workbook(out,{"in_memory":True})
    ws=wb.add_worksheet("source_mailmerge"); evs=wb.add_worksheet("evidence_log")
    hdr=wb.add_format({"bold":True,"bg_color":"#D9EAF7","border":1,"text_wrap":True}); cell=wb.add_format({"border":1,"text_wrap":True})
    for c,h in enumerate(ALL_FIELDS): ws.write(0,c,h,hdr); ws.write(1,c,mapping.get(h,""),cell)
    ws.freeze_panes(1,2); ws.set_column(0,len(ALL_FIELDS)-1,18)
    heads=["field","value","source_document","confidence","note","status_review"]
    for c,h in enumerate(heads):evs.write(0,c,h,hdr)
    for r,f in enumerate(ALL_FIELDS,1):
        e=evidence.get(f,{}); v=mapping.get(f,""); conf=e.get("confidence","")
        try: ok=bool(v) and (conf=="" or float(conf)>=0.75)
        except: ok=False
        status="OK" if ok else ("REVIEW" if v else "MISSING")
        for c,val in enumerate([f,v,e.get("source",""),conf,e.get("note",""),status]):evs.write(r,c,val,cell)
    evs.set_column("A:A",32); evs.set_column("B:B",35); evs.set_column("C:C",28); evs.set_column("E:E",60)
    wb.close(); out.seek(0); return out

# ---------- UI ----------
with st.sidebar:
    st.subheader("⚙️ AI")
    use_ai=st.toggle("Gunakan AI untuk melengkapi",value=True)
    st.session_state["api_key"]=st.text_input("OpenAI API key (opsional bila sudah di Secrets)",type="password",value=st.session_state.get("api_key",""))
    model=st.text_input("Model API",value=st.session_state.get("model","gpt-5")); st.session_state["model"]=model
    st.caption("Excel dihitung secara deterministic. AI hanya melengkapi field yang masih kosong.")

t1,t2,t3,t4=st.tabs(["1 · Upload","2 · Review Data","3 · Evidence","4 · Output"])
with t1:
    uploads=st.file_uploader("Upload semua source untuk satu kegiatan",type=["pdf","docx","xlsx","csv","txt"],accept_multiple_files=True)
    template=st.file_uploader("Template Word laporan (opsional)",type=["docx"])
    if st.button("✨ Proses & isi mailmerge",type="primary",disabled=not uploads,use_container_width=True):
        docs=[]; bar=st.progress(0)
        for i,u in enumerate(uploads):
            b=u.getvalue(); ext=Path(u.name).suffix.lower()
            try: txt=extract_uploaded(u)
            except Exception as e: txt=f"[ERROR {e}]"
            docs.append({"name":u.name,"ext":ext,"bytes":b,"text":txt}); bar.progress((i+1)/len(uploads))
        fields,evidence=deterministic_extract(docs)
        mode="deterministic+structured-xlsx"
        if use_ai and get_api_key():
            try: fields,evidence=ai_enrich(fields,evidence,docs,model); mode+=" + AI"
            except Exception as e: st.warning(f"AI gagal, hasil deterministic tetap dipakai: {e}")
        st.session_state["fields"]=fields; st.session_state["evidence"]=evidence; st.session_state["template_bytes"]=template.getvalue() if template else None
        filled=sum(bool(v) for v in fields.values())
        st.success(f"Selesai · {filled} field mailmerge terisi · mode {mode}. Buka tab Review Data.")
        st.rerun()
with t2:
    fields=st.session_state.get("fields")
    if not fields: st.info("Belum ada hasil. Proses dokumen dulu di tab Upload.")
    else:
        only_filled=st.toggle("Tampilkan hanya field terisi",value=False)
        keys=[k for k in ALL_FIELDS if (fields.get(k) if only_filled else True)]
        rows=[{"Field":k,"Nilai":fields.get(k,"")} for k in keys]
        edited=st.data_editor(rows,hide_index=True,use_container_width=True,num_rows="fixed",column_config={"Field":st.column_config.TextColumn(disabled=True),"Nilai":st.column_config.TextColumn(width="large")})
        for row in edited: fields[row["Field"]]=row["Nilai"]
        st.session_state["fields"]=fields
        st.caption(f"Terisi {sum(bool(v) for v in fields.values())} dari {len(ALL_FIELDS)} field source mailmerge.")
with t3:
    fields=st.session_state.get("fields",{}); evidence=st.session_state.get("evidence",{})
    if not fields: st.info("Belum ada evidence.")
    else:
        rows=[]
        for f in ALL_FIELDS:
            if fields.get(f):
                e=evidence.get(f,{})
                rows.append({"Field":f,"Nilai":fields.get(f,""),"Sumber":e.get("source",""),"Confidence":e.get("confidence",""),"Catatan":e.get("note","")})
        st.dataframe(rows,hide_index=True,use_container_width=True)
with t4:
    fields=st.session_state.get("fields")
    if not fields: st.info("Belum ada output.")
    else:
        evidence=st.session_state.get("evidence",{}); template_bytes=st.session_state.get("template_bytes")
        x=build_xlsx(fields,evidence); d=build_docx(fields,template_bytes)
        safe=re.sub(r"[^A-Za-z0-9_-]+","_",fields.get("nama_pelatihan","laporan"))[:50]
        c1,c2=st.columns(2)
        c1.download_button("⬇️ Source Mailmerge XLSX",x.getvalue(),file_name=f"Source_Mailmerge_{safe}.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
        c2.download_button("⬇️ Laporan Word",d.getvalue(),file_name=f"Laporan_{safe}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
